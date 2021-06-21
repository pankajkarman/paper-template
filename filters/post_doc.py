#!/usr/bin/env python
import re, argparse
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT

def align_table(doc):
    center = WD_TABLE_ALIGNMENT.CENTER
    for table in doc.tables:
        table.alignment = center
    return doc


def make_caption_bold(doc, tag="Figure"):
    for p in doc.paragraphs:
        if tag in p.text:
            for run in p.runs:
                if tag in run.text.split(" ")[0]:
                    texts = run.text.split(":")                    
                    new1 = p.add_run(texts[0] + ":", style=run.style.name)
                    new2 = p.add_run(texts[1:], style=run.style.name)
                    
                    new1.font.bold = True
                    new1.italic = run.italic
                    
                    new2.font.bold = False
                    new2.italic = run.italic  
                    
                    run.text = ""                  
                    break
    return doc

parser = argparse.ArgumentParser(description='DOCX Post-processing.')
parser.add_argument('--infile', help='Give input MS Word filename with docx extension.')
parser.add_argument('--outfile', default='new.docx', help='Give output filename.')
args = parser.parse_args()
print("Arguments:", args)

doc = Document(args.infile)
doc = align_table(doc)
doc = make_caption_bold(doc, tag="Table")
doc = make_caption_bold(doc, tag="Figure")
doc.save(args.outfile)

